from docx import Document

doc = Document()
doc.add_heading('KIRAN KUMAR R', level=1)
doc.add_paragraph('Python Backend Developer')
doc.add_paragraph('Bengaluru, India | +91-7338179551')
doc.add_paragraph('Email: kirankumar@example.com | LinkedIn/GitHub: Kiran')

doc.add_heading('WORK EXPERIENCE', level=2)
doc.add_heading('Software Engineer | CGI (2021 – Present)', level=3)
points = [
    "Built scalable backend services using Python, Django, DRF, improving API response times by 20%.",
    "Designed and implemented RESTful APIs for enterprise systems ensuring smooth data flow.",
    "Worked with microservices architecture, integrating Kafka for asynchronous messaging.",
    "Optimized PostgreSQL/MySQL queries, reducing latency and improving system throughput.",
    "Containerized backend services using Docker and supported deployments on Kubernetes.",
    "Contributed to CI/CD pipelines, code reviews, and Agile processes."
]
for p in points:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('PROJECTS', level=2)
doc.add_heading('URL Shortener Service (Django, DRF, Redis, AWS EC2)', level=3)
points = [
    "Developed a scalable URL shortener supporting high-traffic redirects.",
    "Implemented JWT authentication, RBAC, and Redis caching, improving latency by ~40%.",
    "Normalized database to 3NF and deployed using Nginx + Gunicorn on AWS EC2."
]
for p in points:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('Task Management System (Enterprise Internal Tool)', level=3)
points = [
    "Built a task management backend using DRF with complete CRUD features.",
    "Improved team efficiency by 30% through workflow automation and email notifications.",
    "Added JWT auth, role-based access control, indexing, and optimized database queries."
]
for p in points:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('SKILLS', level=2)
doc.add_paragraph("Languages: Python, SQL")
doc.add_paragraph("Frameworks: Django, Django Rest Framework")
doc.add_paragraph("Databases: PostgreSQL, MySQL")
doc.add_paragraph("Tools: Docker, Kubernetes, Redis, Kafka, Git")
doc.add_paragraph("Cloud: AWS (EC2, S3, IAM), Basic CI/CD")
doc.add_paragraph("Concepts: Microservices, REST API Design, Caching, DB Normalization, System Design")

doc.add_heading('EDUCATION', level=2)
doc.add_paragraph("CMR University — B.Tech in Computer Science & Engineering (2018–2022)")

doc.add_heading('ACHIEVEMENTS', level=2)
points = [
    "Delivered backend modules used across multiple teams at CGI.",
    "Enhanced DB performance using indexing & normalization techniques.",
    "Recognized for delivering high-quality backend components consistently."
]
for p in points:
    doc.add_paragraph(p, style='List Bullet')

filename = "Kiran_Kumar_Backend_Developer_Resume.docx"
doc.save(filename)

print("Resume created:", filename)
